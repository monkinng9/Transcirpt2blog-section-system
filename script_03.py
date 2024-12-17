from pathlib import Path
import sys
from config import INPUT_MARKDOWN, OUTPUT_MARKDOWN
from docx import Document
from docx.shared import Inches
import markdown
from bs4 import BeautifulSoup
import re

def convert_markdown_to_word(input_file, output_file):
    """
    Convert a markdown file to Word document (.docx)
    """
    # Read markdown content
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Convert markdown to HTML
    html_content = markdown.markdown(md_content)
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create Word document
    doc = Document()

    # Process HTML elements
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'img']):
        if element.name.startswith('h'):
            level = int(element.name[1])
            doc.add_heading(element.text, level=level)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'img':
            # Handle images
            src = element.get('src')
            if src and Path(src).exists():
                doc.add_picture(src, width=Inches(6))
                # Add caption if alt text exists
                if element.get('alt'):
                    doc.add_paragraph(element.get('alt')).italic = True

    # Save the document
    doc.save(output_file)
    print(f"Successfully converted {input_file} to {output_file}")

def main():
    input_file = INPUT_MARKDOWN
    output_file = Path(input_file).stem + '.docx'

    if not Path(input_file).exists():
        print(f"Error: Input file {input_file} does not exist")
        sys.exit(1)

    convert_markdown_to_word(input_file, output_file)

if __name__ == "__main__":
    main()
